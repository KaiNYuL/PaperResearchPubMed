from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from .config import DATA_DIR


def _render_markdown(papers: List[Dict[str, Any]]) -> str:
    lines = ["# 论文整理结果", ""]
    for idx, paper in enumerate(papers, 1):
        lines.append(f"## {idx}. {paper.get('title', '')}")
        lines.append(f"- 关键字：{', '.join(paper.get('keywords', []))}")
        lines.append(f"- 匹配维度：{paper.get('match_dimension', '')}")
        lines.append(f"- 发表时间：{paper.get('published_date', '')}")
        if paper.get("doi"):
            lines.append(f"- DOI：{paper.get('doi')}")
        if paper.get("authors"):
            lines.append(f"- 作者：{', '.join(paper.get('authors', []))}")
        lines.append("- 摘要：")
        lines.append(paper.get("abstract", ""))
        lines.append("")
    return "\n".join(lines)


def _render_txt(papers: List[Dict[str, Any]]) -> str:
    lines = []
    for idx, paper in enumerate(papers, 1):
        lines.append(f"{idx}. {paper.get('title', '')}")
        lines.append(f"关键字：{', '.join(paper.get('keywords', []))}")
        lines.append(f"匹配维度：{paper.get('match_dimension', '')}")
        lines.append(f"发表时间：{paper.get('published_date', '')}")
        if paper.get("doi"):
            lines.append(f"DOI：{paper.get('doi')}")
        if paper.get("authors"):
            lines.append(f"作者：{', '.join(paper.get('authors', []))}")
        lines.append("摘要：")
        lines.append(paper.get("abstract", ""))
        lines.append("-" * 40)
    return "\n".join(lines)


def export_docs(papers: List[Dict[str, Any]], fmt: str, output_dir: str | None, filename: str | None) -> str:
    fmt = (fmt or "markdown").lower()
    default_dir = DATA_DIR / "exports"
    output_path = Path(output_dir) if output_dir else default_dir
    output_path.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if fmt in {"md", "markdown"}:
        name = filename or f"papers_{stamp}.md"
        content = _render_markdown(papers)
        path = output_path / name
        path.write_text(content, encoding="utf-8")
        return str(path)
    if fmt in {"txt", "text"}:
        name = filename or f"papers_{stamp}.txt"
        content = _render_txt(papers)
        path = output_path / name
        path.write_text(content, encoding="utf-8")
        return str(path)
    if fmt in {"doc", "docx", "word"}:
        try:
            from docx import Document
        except Exception as exc:
            raise ValueError(f"Word 导出不可用，请检查 lxml/python-docx 安装：{exc}")
        name = filename or f"papers_{stamp}.docx"
        path = output_path / name
        doc = Document()
        doc.add_heading("论文整理结果", level=1)
        for idx, paper in enumerate(papers, 1):
            doc.add_heading(f"{idx}. {paper.get('title', '')}", level=2)
            doc.add_paragraph(f"关键字：{', '.join(paper.get('keywords', []))}")
            doc.add_paragraph(f"匹配维度：{paper.get('match_dimension', '')}")
            doc.add_paragraph(f"发表时间：{paper.get('published_date', '')}")
            if paper.get("doi"):
                doc.add_paragraph(f"DOI：{paper.get('doi')}")
            if paper.get("authors"):
                doc.add_paragraph(f"作者：{', '.join(paper.get('authors', []))}")
            doc.add_paragraph("摘要：")
            doc.add_paragraph(paper.get("abstract", ""))
        doc.save(path)
        return str(path)
    raise ValueError("不支持的导出格式")
